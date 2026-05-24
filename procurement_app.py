import os
import io
import json
import re

import streamlit as st
import pandas as pd
import pdfplumber
from langchain_openai import ChatOpenAI
from docx import Document
from openpyxl import load_workbook

# ─── Page Config ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Procurement Comparison Assistant",
    page_icon="🔍",
    layout="wide",
)

# ─── Styling ──────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.main { background: #f8fafc; }
.hero-header {
    background: linear-gradient(135deg, #1e3a5f 0%, #2563eb 60%, #3b82f6 100%);
    border-radius: 16px; padding: 36px 40px; margin-bottom: 28px;
    color: white; box-shadow: 0 8px 32px rgba(37,99,235,0.18);
}
.hero-header h1 { font-size: 2rem; font-weight: 700; margin: 0 0 6px; }
.hero-header p  { font-size: 1rem; opacity: 0.85; margin: 0; }
.step-indicator {
    display: flex; align-items: center; gap: 10px;
    background: white; border-radius: 10px;
    padding: 14px 18px; margin-bottom: 14px;
    box-shadow: 0 1px 6px rgba(0,0,0,0.06);
    font-size: .9rem; color: #334155;
}
.step-num {
    background: #2563eb; color: white; border-radius: 50%;
    width: 26px; height: 26px; display: flex; align-items: center;
    justify-content: center; font-weight: 700; font-size: .8rem; flex-shrink: 0;
}
.metric-box {
    background: white; border-radius: 12px; padding: 18px 20px;
    text-align: center; box-shadow: 0 1px 8px rgba(0,0,0,0.07);
}
.metric-num { font-size: 2rem; font-weight: 700; color: #1e3a5f; }
.metric-lbl { font-size: .78rem; color: #64748b; text-transform: uppercase;
              letter-spacing: .05em; margin-top: 2px; }
.badge-identical { background:#dcfce7; color:#166534; padding:4px 14px;
    border-radius:999px; font-size:.78rem; font-weight:600; display:inline-block; }
.badge-diff      { background:#fef3c7; color:#92400e; padding:4px 14px;
    border-radius:999px; font-size:.78rem; font-weight:600; display:inline-block; }
.badge-missing   { background:#fee2e2; color:#991b1b; padding:4px 14px;
    border-radius:999px; font-size:.78rem; font-weight:600; display:inline-block; }
.diff-row {
    background:#fffbeb; border-left:4px solid #f59e0b;
    border-radius:0 8px 8px 0; padding:14px 18px; margin:10px 0;
}
.diff-label { font-weight:600; font-size:.82rem; color:#78350f; margin-bottom:4px; }
.diff-val   { font-size:.88rem; color:#451a03; }
.ok-row {
    background:#f0fdf4; border-left:4px solid #22c55e;
    border-radius:0 8px 8px 0; padding:10px 18px; margin:8px 0;
    font-size:.88rem; color:#14532d;
}
.summary-box {
    background:linear-gradient(135deg,#eff6ff,#dbeafe);
    border:1px solid #bfdbfe; border-radius:12px; padding:22px 26px; margin-bottom:22px;
}
.summary-box h3 { color:#1d4ed8; margin:0 0 10px; font-size:1rem; }
.stButton > button {
    background: linear-gradient(135deg, #1e3a5f, #2563eb) !important;
    color: white !important; border: none !important;
    border-radius: 10px !important; padding: 12px 32px !important;
    font-weight: 600 !important; font-size: 1rem !important;
    width: 100%; transition: opacity .2s !important;
}
.stButton > button:hover { opacity: .88 !important; }
</style>
""", unsafe_allow_html=True)


# ─── Models ───────────────────────────────────────────────────────────────────

MODELS = {
    "Qwen3.6 Plus":                       "qwen/qwen3.6-plus",
}

def get_llm(model_id: str) -> ChatOpenAI:

    api_key = os.getenv("OPENROUTER_API_KEY")

    if not api_key:
        st.error("OPENROUTER_API_KEY environment variable is missing.")
        st.stop()

    return ChatOpenAI(
        openai_api_key=api_key,
        openai_api_base="https://openrouter.ai/api/v1",
        model=model_id,
        temperature=0,
        max_tokens=8192,
    )



# ─── Text Extractors ──────────────────────────────────────────────────────────

def extract_pdf(file_bytes: bytes) -> str:
    """pdfplumber OCR — pulls text and tables from every page."""
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            parts.append(f"[Page {i}]")
            # Tables first — preserves column structure
            for table in page.extract_tables():
                for row in table:
                    parts.append("\t".join(str(c) if c else "" for c in row))
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts)


def extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    return "\n".join(parts)


def extract_xlsx(file_bytes: bytes) -> str:
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"=== Sheet: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                parts.append("\t".join(str(c) if c is not None else "" for c in row))
    return "\n".join(parts)


def extract_csv(file_bytes: bytes) -> str:
    try:
        return pd.read_csv(io.BytesIO(file_bytes)).to_csv(index=False)
    except Exception:
        return file_bytes.decode("utf-8", errors="replace")


def read_file(uploaded_file) -> str:
    """Return plain text for any supported file type."""
    name = uploaded_file.name.lower()
    raw  = uploaded_file.read()
    if name.endswith(".pdf"):
        return extract_pdf(raw)
    elif name.endswith(".docx"):
        return extract_docx(raw)
    elif name.endswith((".xlsx", ".xlsm", ".xls")):
        return extract_xlsx(raw)
    elif name.endswith((".csv", ".tsv")):
        return extract_csv(raw)
    else:
        return raw.decode("utf-8", errors="replace")


# ─── Prompt & LLM Call ────────────────────────────────────────────────────────

PROMPT = """You are a senior procurement analyst.

Below are the full text contents of two product documents (Document A and Document B), extracted via OCR.

Read both carefully, understand them in full, then compare them directly.

Output ONLY a raw JSON object — no markdown fences, no explanation, no preamble.

JSON schema:
{
  "doc_a_product_count": <int>,
  "doc_b_product_count": <int>,
  "matched_count": <int>,
  "identical_count": <int>,
  "different_count": <int>,
  "only_in_a_count": <int>,
  "only_in_b_count": <int>,
  "overall_status": "IDENTICAL" | "DIFFERENCES FOUND",
  "products": [
    {
      "id": "<SKU / Part# / ID exactly as written>",
      "status": "IDENTICAL" | "DIFFERENT" | "ONLY_IN_A" | "ONLY_IN_B",
      "doc_a": {"name": "...", "description": "..."},
      "doc_b": {"name": "...", "description": "..."},
      "differences": [
        {"field": "name | description | other", "doc_a_value": "...", "doc_b_value": "..."}
      ]
    }
  ],
}

Rules:
- Match products across documents by their identifier (SKU, Part#, ID, item code — whatever label is used).
- Keep all IDs exactly as written — never reformat or normalize them.
- ONLY_IN_A → doc_b = null, differences = [].
- ONLY_IN_B → doc_a = null, differences = [].
- IDENTICAL → differences = [].
- Make sure of the missing values in the JSON file are null not empty string or any other value.
- Even minor wording differences count as DIFFERENT.
- Never hallucinate products not present in the documents.
"""


def parse_llm_response(response) -> str:
    text = response.content
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)
    text = re.sub(r"```[a-z]*\n?", "", text)
    return text.strip().rstrip("`").strip()


def parse_json(raw: str) -> dict:
    match = re.search(r'\{.*\}', raw, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in model response.")
    return json.loads(match.group())


def compare(text_a: str, text_b: str, llm: ChatOpenAI) -> dict:
    msg = f"{PROMPT}\n\n--- DOCUMENT A ---\n{text_a}\n\n--- DOCUMENT B ---\n{text_b}"
    response = llm.invoke([{"role": "user", "content": msg}])
    return parse_json(parse_llm_response(response))


# ─── Results Renderer ─────────────────────────────────────────────────────────

def render_results(result: dict):
    st.markdown("---")
    c1, c2, c3, c4, c5 = st.columns(5)
    for col, (num, lbl, color) in zip([c1,c2,c3,c4,c5], [
        (result.get('doc_a_product_count','?'), "Doc A",         None),
        (result.get('doc_b_product_count','?'), "Doc B",         None),
        (result.get('identical_count','?'),     "Identical",     "#22c55e"),
        (result.get('different_count','?'),     "Different",     "#f59e0b"),
        (result.get('only_in_a_count',0)+result.get('only_in_b_count',0), "Missing/Extra", "#ef4444"),
    ]):
        clr = f'style="color:{color}"' if color else ""
        col.markdown(f'<div class="metric-box"><div class="metric-num" {clr}>{num}</div>'
                     f'<div class="metric-lbl">{lbl}</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    if result.get("overall_status") == "IDENTICAL":
        st.success("✅ **RESULT: IDENTICAL** — Both documents contain the same products with matching names and descriptions.")
    else:
        n = result.get("different_count",0) + result.get("only_in_a_count",0) + result.get("only_in_b_count",0)
        st.warning(f"⚠️ **RESULT: DIFFERENCES FOUND** — {n} discrepanc{'y' if n==1 else 'ies'} detected.")


    products = result.get("products", [])
    if not products:
        return

    st.markdown("### Product-by-Product Breakdown")
    counts = {s: sum(1 for p in products if p.get("status")==s)
              for s in ("IDENTICAL","DIFFERENT","ONLY_IN_A","ONLY_IN_B")}
    missing = counts["ONLY_IN_A"] + counts["ONLY_IN_B"]

    tab_all, tab_diff, tab_ok, tab_miss = st.tabs([
        f"All ({len(products)})",
        f"🟡 Different ({counts['DIFFERENT']})",
        f"🟢 Identical ({counts['IDENTICAL']})",
        f"🔴 Missing/Extra ({missing})",
    ])

    def render_product(p):
        pid     = p.get("id", "?")
        status  = p.get("status", "?")
        doc_a   = p.get("doc_a") or {}
        doc_b   = p.get("doc_b") or {}
        diffs   = p.get("differences", [])
        badge   = {"IDENTICAL":"badge-identical","DIFFERENT":"badge-diff",
                   "ONLY_IN_A":"badge-missing","ONLY_IN_B":"badge-missing"}.get(status,"badge-diff")
        label   = {"IDENTICAL":"✓ Identical",
                   "DIFFERENT":f"⚠ {len(diffs)} difference{'s' if len(diffs)!=1 else ''}",
                   "ONLY_IN_A":"Only in Doc A","ONLY_IN_B":"Only in Doc B"}.get(status, status)
        title   = doc_a.get("name") or doc_b.get("name") or ""

        with st.expander(f"**{pid}** — {title}", expanded=(status=="DIFFERENT")):
            st.markdown(f'<span class="{badge}">{label}</span>', unsafe_allow_html=True)
            st.markdown("")
            if status == "IDENTICAL":
                st.markdown(f'<div class="ok-row"><strong>Name:</strong> {doc_a.get("name","—")}<br>'
                            f'<strong>Description:</strong> {doc_a.get("description","—")}</div>',
                            unsafe_allow_html=True)
            elif status == "DIFFERENT":

                st.markdown("**Differences:**")
                for d in diffs:
                    st.markdown(f'<div class="diff-row"><div class="diff-label">{d.get("field","?").upper()}</div>'
                                f'<div class="diff-val"><span style="color:#dc2626">A: {d.get("doc_a_value","—")}</span><br>'
                                f'<span style="color:#16a34a">B: {d.get("doc_b_value","—")}</span></div></div>',
                                unsafe_allow_html=True)
            elif status == "ONLY_IN_A":
                st.markdown(f"**Only in Document A**  \n**Name:** {doc_a.get('name','—')}  \n**Desc:** {doc_a.get('description','—')}")
            elif status == "ONLY_IN_B":
                st.markdown(f"**Only in Document B**  \n**Name:** {doc_b.get('name','—')}  \n**Desc:** {doc_b.get('description','—')}")

    with tab_all:
        for p in products: render_product(p)
    with tab_diff:
        items = [p for p in products if p.get("status")=="DIFFERENT"]
        [render_product(p) for p in items] if items else st.success("No differences!")
    with tab_ok:
        items = [p for p in products if p.get("status")=="IDENTICAL"]
        [render_product(p) for p in items] if items else st.info("No identical products.")
    with tab_miss:
        items = [p for p in products if p.get("status") in ("ONLY_IN_A","ONLY_IN_B")]
        [render_product(p) for p in items] if items else st.success("No missing or extra products.")

    st.markdown("---")
    st.download_button("⬇️ Download Report (JSON)",
                       data=json.dumps(result, indent=2, ensure_ascii=False),
                       file_name="procurement_comparison.json", mime="application/json")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    st.markdown("""
    <div class="hero-header">
        <h1>🔍 Procurement Comparison Assistant</h1>
        <p>Upload two product documents — PDF, Word, or Excel — and get an AI-powered comparison of SKUs, names, and descriptions.</p>
    </div>
    """, unsafe_allow_html=True)

    with st.sidebar:
        st.markdown("### ⚙️ Settings")
        api_key = os.getenv("OPENROUTER_API_KEY")

 
        st.markdown("---")
        st.markdown("📁 Supported Formats")
        st.markdown("- PDF `.pdf` *(pdfplumber OCR)*\n- Word `.docx`\n- Excel `.xlsx / .xls`\n- CSV `.csv`")
        st.markdown("---")


    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown('<div class="step-indicator"><div class="step-num">1</div>Upload Document A</div>',
                    unsafe_allow_html=True)
    with c2:
        st.markdown('<div class="step-indicator"><div class="step-num">2</div>Upload Document B</div>',
                    unsafe_allow_html=True)
    with c3:
        st.markdown('<div class="step-indicator"><div class="step-num">3</div>Click Compare</div>',
                    unsafe_allow_html=True)

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("**📄 Document A**")
        file_a = st.file_uploader("Upload Document A", type=["pdf","docx","xlsx","xlsm","xls","csv"],
                                   key="file_a", label_visibility="collapsed")
        if file_a: st.success(f"✓ {file_a.name}")

    with col_b:
        st.markdown("**📄 Document B**")
        file_b = st.file_uploader("Upload Document B", type=["pdf","docx","xlsx","xlsm","xls","csv"],
                                   key="file_b", label_visibility="collapsed")
        if file_b: st.success(f"✓ {file_b.name}")

    st.markdown("<br>", unsafe_allow_html=True)
    run = st.button("🔍 Compare Documents", disabled=(not file_a or not file_b))

    if run and file_a and file_b and api_key:
        with st.spinner("Reading documents…"):
            text_a = read_file(file_a)
            text_b = read_file(file_b)

        col1, col2 = st.columns(2)

        with st.spinner("Comparing documents with AI… This may take a moment."):
            llm = get_llm("qwen/qwen3.6-plus")
            result = compare(text_a, text_b, llm)

        render_results(result)


if __name__ == "__main__":
    main()
